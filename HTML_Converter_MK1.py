"""Преобразования Word документа в текст с HTML тегами (HTML формат)"""

#pip install python-docx
import docx



file_1=str(input("Введите имя файла:"))


doc = docx.Document(file_1)                 #Блок вытаскивает текст из файла.
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)

kol=(len(doc.paragraphs))           #Подсчёт количества абзацев.

with open("new.txt", "a") as f:             #Ставит начальный тег.
    f.write("<p>")
    f.close()

for i in range(kol):                      #Находи и заменяем всё необходимое на теги. Начало:
    rep=str(text[i])
    if "(" in rep:
        rep=rep.replace("(","<em>(")
    if ")" in rep:
        rep=rep.replace(")",")</em>")

    if "«" in rep:
        rep=rep.replace("«","<em><strong>«")
    if "»" in rep:
        rep=rep.replace("»","»</strong></em>")

    if "“" in rep:
        rep=rep.replace("“","<em><strong>“")
    if "”" in rep:
        rep=rep.replace("”","”</strong></em>")

    if ":" in rep:
        ooo=rep.split()
        kol_1 = (len(ooo))
        rep_dv = ""
        for i2 in range(kol_1):
            t = str(ooo[i2])
            if ":" in t:
                rep_dv = rep_dv + "<strong>"+t+"</strong> "
            else:
                u=ooo[i2]
                rep_dv = rep_dv + u +" "
        rep=rep_dv

    if "–" in rep:
        ooo=rep.split()
        kol_1 = (len(ooo))
        rep_dv = ""

        for i2 in range(kol_1):
            isk= str(ooo[i2])
            i3=i2+1
            if i3>=kol_1:
                i3=i3-1
            t = str(ooo[i3])
            if "–" in t:
                rep_dv = rep_dv + "<strong>"+isk+"</strong> "
            else:
                u=ooo[i2]
                rep_dv = rep_dv + u +" "
        rep=rep_dv

    # Конец

    with open("new.txt", "a") as f:              #Записывае полученный абзац в новый текстовый фаил.
        f.write(rep+"</p>"+"\n")
        f.close()











#MK2  -  наброски для второй версии.

"""bold — Полужирное начертание
underline — Подчеркнутый текст
italic — Курсивное начертание
strike — Зачеркнутый текст



for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        print('Полужирный текст:', run.bold)
        print('Подчёркнутый текст:', run.underline)
        print('курсивный текст:', run.italic)
        print("   ")
"""
